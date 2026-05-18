

import base64
import os
import re
import smtplib
import uuid
from datetime import datetime, timezone
from email.message import EmailMessage
from pathlib import Path
from typing import Any, Optional

from dotenv import load_dotenv

load_dotenv()


# ============================================================
# PATHS
# ============================================================

OUTPUT_DIR = Path(os.getenv("OUTPUT_DIR", "data/outputs"))
VECTORSTORE_DIR = Path(os.getenv("VECTORSTORE_DIR", "data/vectorstore"))
RAW_PDFS_DIR = Path(os.getenv("RAW_PDFS_DIR", "data/raw_pdfs"))

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
VECTORSTORE_DIR.mkdir(parents=True, exist_ok=True)
RAW_PDFS_DIR.mkdir(parents=True, exist_ok=True)


# ============================================================
# BASIC HELPERS
# ============================================================

def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def new_id(prefix: str) -> str:
    return f"{prefix}_{uuid.uuid4().hex[:10]}"


def safe_filename(name: str, default: str = "artifact") -> str:
    """
    Create a safe filename from arbitrary user/model text.
    """
    name = name.strip() or default
    name = re.sub(r"[^\w\-. ]+", "", name, flags=re.UNICODE)
    name = re.sub(r"\s+", "_", name)
    return name[:120] or default


def short_preview(content: str, limit: int = 500) -> str:
    content = content or ""
    content = content.strip()
    if len(content) <= limit:
        return content
    return content[:limit].rstrip() + "..."


def find_email_in_text(text: str) -> Optional[str]:
    if not text:
        return None

    match = re.search(r"[\w.\-+%]+@[\w.\-]+\.[A-Za-z]{2,}", text)
    if not match:
        return None

    return match.group(0).strip(".,;:()[]{}<>")


def is_yes_confirmation(text: str) -> bool:
    """
    Detect short confirmations in Arabic/English.
    This is intentionally simple; the orchestrator should also use pending_approval.
    """
    text = (text or "").strip().lower()

    yes_phrases = {
        "yes",
        "y",
        "send",
        "send it",
        "confirm",
        "confirmed",
        "ok",
        "okay",
        "sure",
        "go ahead",
        "do it",
        "تمام",
        "اوك",
        "أوك",
        "اه",
        "آه",
        "ايوه",
        "أيوه",
        "ابعت",
        "ابعته",
        "ابعت دلوقتي",
        "ابعته دلوقتي",
        "موافق",
        "أكد",
        "اكيد",
    }

    return text in yes_phrases


# ============================================================
# IMAGE HELPERS
# ============================================================

def guess_mime_type(path: Path) -> str:
    suffix = path.suffix.lower()

    if suffix in [".jpg", ".jpeg"]:
        return "image/jpeg"
    if suffix == ".webp":
        return "image/webp"
    if suffix == ".gif":
        return "image/gif"

    return "image/png"


def image_to_data_url(image_path: str) -> str:
    """
    Convert a local image file into a data URL for vision models.
    """
    path = Path(image_path)

    if not path.exists():
        raise FileNotFoundError(f"Image not found: {image_path}")

    mime = guess_mime_type(path)
    encoded = base64.b64encode(path.read_bytes()).decode("utf-8")
    return f"data:{mime};base64,{encoded}"


# ============================================================
# WEB SEARCH
# ============================================================

def tavily_search(query: str, max_results: int = 5) -> list[dict[str, Any]]:
    """
    Search the web using Tavily.

    Returns a normalized list. If Tavily is not configured, returns [].
    """
    api_key = os.getenv("TAVILY_API_KEY")

    if not api_key:
        return []

    try:
        from tavily import TavilyClient

        client = TavilyClient(api_key=api_key)
        result = client.search(
            query=query,
            max_results=max_results,
            include_answer=False,
            include_raw_content=False,
        )

        return result.get("results", []) or []

    except Exception as exc:
        return [
            {
                "title": "Tavily search error",
                "url": "",
                "content": str(exc),
                "score": 0.0,
                "error": True,
            }
        ]


def register_sources_from_web(results: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """
    Convert raw Tavily results into source items used by the graph state.
    """
    sources: list[dict[str, Any]] = []

    for item in results:
        content = item.get("content") or item.get("snippet") or ""

        sources.append(
            {
                "source_id": new_id("source"),
                "type": "web",
                "title": item.get("title", ""),
                "url": item.get("url", ""),
                "content": content,
                "snippet": short_preview(content, 350),
                "score": item.get("score"),
                "metadata": {
                    "created_at": utc_now_iso(),
                    "raw": item,
                },
            }
        )

    return sources


def format_sources_for_prompt(sources: list[dict[str, Any]], limit: int = 8) -> str:
    """
    Convert sources into a compact prompt-friendly format.
    """
    if not sources:
        return "No sources available."

    chunks: list[str] = []

    for idx, source in enumerate(sources[:limit], start=1):
        title = source.get("title") or "Untitled"
        url = source.get("url") or source.get("path") or ""
        content = source.get("content") or source.get("snippet") or ""

        chunks.append(
            f"[{idx}] {title}\n"
            f"URL/Path: {url}\n"
            f"Content: {short_preview(content, 1200)}"
        )

    return "\n\n".join(chunks)


# ============================================================
# ARTIFACTS / FILES
# ============================================================

def make_artifact(
    *,
    artifact_type: str,
    path: Path,
    title: Optional[str] = None,
    content_preview: str = "",
    created_from: str = "",
    status: str = "ready",
    metadata: Optional[dict[str, Any]] = None,
    available_actions: Optional[list[str]] = None,
) -> dict[str, Any]:
    artifact_id = new_id("artifact")

    return {
        "artifact_id": artifact_id,
        "type": artifact_type,
        "title": title or path.stem,
        "filename": path.name,
        "path": str(path),
        "content_preview": short_preview(content_preview, 500),
        "status": status,
        "created_from": created_from,
        "available_actions": available_actions or ["edit", "save", "email"],
        "metadata": {
            "created_at": utc_now_iso(),
            **(metadata or {}),
        },
    }


def save_text_file(
    content: str,
    filename: Optional[str] = None,
    *,
    title: Optional[str] = None,
    artifact_type: str = "text",
    created_from: str = "",
) -> dict[str, Any]:
    artifact_id = new_id("file")

    if filename is None:
        filename = f"{artifact_id}.txt"

    filename = safe_filename(filename)
    path = OUTPUT_DIR / filename

    path.write_text(content or "", encoding="utf-8")

    return make_artifact(
        artifact_type=artifact_type,
        path=path,
        title=title,
        content_preview=content or "",
        created_from=created_from,
        available_actions=["edit", "save", "email"],
    )


def save_markdown(
    content: str,
    filename: Optional[str] = None,
    *,
    title: Optional[str] = None,
    artifact_type: str = "markdown",
    created_from: str = "",
) -> dict[str, Any]:
    artifact_id = new_id("md")

    if filename is None:
        filename = f"{artifact_id}.md"

    if not filename.lower().endswith(".md"):
        filename += ".md"

    filename = safe_filename(filename)
    path = OUTPUT_DIR / filename

    path.write_text(content or "", encoding="utf-8")

    return make_artifact(
        artifact_type=artifact_type,
        path=path,
        title=title,
        content_preview=content or "",
        created_from=created_from,
        available_actions=["edit", "export_pdf", "email"],
    )


def export_pdf_from_markdown(
    markdown_content: str,
    filename: Optional[str] = None,
    *,
    title: Optional[str] = None,
    created_from: str = "",
) -> dict[str, Any]:
    """
    Export Markdown to PDF when optional deps are available.

    Professional behavior:
    - Try Markdown -> HTML -> PDF via markdown + weasyprint.
    - If unavailable or PDF render fails, save Markdown fallback.
    """
    artifact_id = new_id("pdf")

    if filename is None:
        filename = f"{artifact_id}.pdf"

    if not filename.lower().endswith(".pdf"):
        filename += ".pdf"

    filename = safe_filename(filename)
    pdf_path = OUTPUT_DIR / filename

    try:
        import markdown as md
        from weasyprint import HTML

        html_body = md.markdown(
            markdown_content or "",
            extensions=["extra", "tables", "fenced_code"],
        )

        html = f"""
        <!doctype html>
        <html>
        <head>
          <meta charset="utf-8">
          <style>
            body {{
              font-family: Arial, sans-serif;
              line-height: 1.6;
              margin: 36px;
            }}
            code, pre {{
              font-family: Consolas, monospace;
            }}
            pre {{
              background: #f5f5f5;
              padding: 12px;
              border-radius: 8px;
              overflow-wrap: break-word;
              white-space: pre-wrap;
            }}
            table {{
              border-collapse: collapse;
              width: 100%;
            }}
            th, td {{
              border: 1px solid #ddd;
              padding: 8px;
            }}
          </style>
        </head>
        <body>
        {html_body}
        </body>
        </html>
        """

        HTML(string=html).write_pdf(str(pdf_path))

        return make_artifact(
            artifact_type="pdf",
            path=pdf_path,
            title=title,
            content_preview=markdown_content or "",
            created_from=created_from,
            available_actions=["save", "email"],
            metadata={"rendered_from": "markdown"},
        )

    except Exception as exc:
        # Fallback 1: try ReportLab so Windows/dev machines still get a real PDF
        # even when WeasyPrint native dependencies are unavailable.
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import cm
            from reportlab.pdfgen import canvas

            c = canvas.Canvas(str(pdf_path), pagesize=A4)
            width, height = A4
            left = 2 * cm
            top = height - 2 * cm
            bottom = 2 * cm
            line_height = 14

            plain_text = re.sub(r"[`*_>#]+", "", markdown_content or "")
            wrapped_lines: list[str] = []

            for paragraph in plain_text.splitlines():
                paragraph = paragraph.rstrip()
                if not paragraph:
                    wrapped_lines.append("")
                    continue

                while len(paragraph) > 95:
                    cut = paragraph.rfind(" ", 0, 95)
                    if cut <= 20:
                        cut = 95
                    wrapped_lines.append(paragraph[:cut].strip())
                    paragraph = paragraph[cut:].strip()
                wrapped_lines.append(paragraph)

            text_obj = c.beginText(left, top)
            text_obj.setFont("Helvetica", 10)
            y = top

            for line in wrapped_lines:
                if y < bottom:
                    c.drawText(text_obj)
                    c.showPage()
                    text_obj = c.beginText(left, top)
                    text_obj.setFont("Helvetica", 10)
                    y = top

                text_obj.textLine(line)
                y -= line_height

            c.drawText(text_obj)
            c.save()

            return make_artifact(
                artifact_type="pdf",
                path=pdf_path,
                title=title,
                content_preview=markdown_content or "",
                created_from=created_from,
                available_actions=["save", "email"],
                metadata={
                    "rendered_from": "markdown",
                    "renderer": "reportlab_fallback",
                    "weasyprint_error": str(exc),
                },
            )

        except Exception as reportlab_exc:
            # Fallback 2: save Markdown so the workflow still returns something.
            fallback_name = filename.replace(".pdf", ".md")
            artifact = save_markdown(
                markdown_content,
                filename=fallback_name,
                title=title,
                artifact_type="markdown",
                created_from=created_from,
            )
            artifact["metadata"]["pdf_error"] = str(exc)
            artifact["metadata"]["reportlab_error"] = str(reportlab_exc)
            artifact["metadata"]["note"] = "PDF rendering failed; saved Markdown fallback."
            return artifact



def _markdown_to_plain_lines(markdown_content: str) -> list[str]:
    """
    Lightweight markdown-to-readable-text helper for fallback exporters.
    Keeps content understandable without depending on a full markdown parser.
    """
    text = markdown_content or ""
    text = re.sub(r"```[\s\S]*?```", lambda m: m.group(0).replace("```", ""), text)
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.splitlines()


def export_docx_from_markdown(
    markdown_content: str,
    filename: Optional[str] = None,
    *,
    title: Optional[str] = None,
    created_from: str = "",
) -> dict[str, Any]:
    """
    Export Markdown-ish content to a Word .docx document.

    Requires:
        pip install python-docx

    If python-docx is not available, it saves a .txt fallback and clearly marks it.
    """
    artifact_id = new_id("docx")

    if filename is None:
        filename = f"{artifact_id}.docx"

    if not filename.lower().endswith(".docx"):
        filename += ".docx"

    filename = safe_filename(filename)
    docx_path = OUTPUT_DIR / filename

    try:
        from docx import Document

        document = Document()

        if title:
            document.add_heading(title, level=1)

        in_code_block = False
        code_lines: list[str] = []

        for raw_line in (markdown_content or "").splitlines():
            line = raw_line.rstrip()

            if line.strip().startswith("```"):
                if not in_code_block:
                    in_code_block = True
                    code_lines = []
                else:
                    if code_lines:
                        paragraph = document.add_paragraph()
                        run = paragraph.add_run("\n".join(code_lines))
                        run.font.name = "Consolas"
                    in_code_block = False
                    code_lines = []
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            stripped = line.strip()

            if not stripped:
                document.add_paragraph("")
                continue

            heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
            if heading:
                level = min(len(heading.group(1)), 4)
                document.add_heading(heading.group(2).strip(), level=level)
                continue

            bullet = re.match(r"^[-*]\s+(.+)$", stripped)
            if bullet:
                document.add_paragraph(bullet.group(1).strip(), style="List Bullet")
                continue

            numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
            if numbered:
                document.add_paragraph(numbered.group(1).strip(), style="List Number")
                continue

            cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", stripped)
            cleaned = re.sub(r"`([^`]+)`", r"\1", cleaned)
            document.add_paragraph(cleaned)

        if in_code_block and code_lines:
            paragraph = document.add_paragraph()
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Consolas"

        document.save(str(docx_path))

        return make_artifact(
            artifact_type="docx",
            path=docx_path,
            title=title,
            content_preview=markdown_content or "",
            created_from=created_from,
            available_actions=["save", "email"],
            metadata={"rendered_from": "markdown"},
        )

    except Exception as exc:
        fallback_name = filename.replace(".docx", ".txt")
        artifact = save_text_file(
            markdown_content,
            filename=fallback_name,
            title=title,
            artifact_type="text",
            created_from=created_from,
        )
        artifact["metadata"]["docx_error"] = str(exc)
        artifact["metadata"]["note"] = "DOCX rendering failed; saved text fallback."
        return artifact


def read_artifact_content(artifact: dict[str, Any]) -> str:
    path = artifact.get("path")
    if not path:
        return ""

    p = Path(path)
    if not p.exists() or not p.is_file():
        return ""

    try:
        return p.read_text(encoding="utf-8")
    except Exception:
        return ""


def find_artifact_by_id(artifacts: list[dict[str, Any]], artifact_id: Optional[str]) -> Optional[dict[str, Any]]:
    if not artifact_id:
        return None

    for artifact in artifacts:
        if artifact.get("artifact_id") == artifact_id:
            return artifact

    return None


def latest_artifact(
    artifacts: list[dict[str, Any]],
    *,
    preferred_types: Optional[list[str]] = None,
) -> Optional[dict[str, Any]]:
    if not artifacts:
        return None

    if preferred_types:
        for artifact in reversed(artifacts):
            if artifact.get("type") in preferred_types:
                return artifact

    return artifacts[-1]


# ============================================================
# EMAIL
# ============================================================

def create_email_draft(
    *,
    to: str,
    subject: str,
    body: str,
    attachment_artifact_id: Optional[str] = None,
    attachment_path: Optional[str] = None,
) -> dict[str, Any]:
    """
    Create a draft object. This does not send anything.
    """
    draft_id = new_id("email")

    return {
        "draft_id": draft_id,
        "to": to,
        "subject": subject,
        "body": body,
        "attachment_artifact_id": attachment_artifact_id,
        "attachment_path": attachment_path,
        "status": "draft_ready",
        "requires_confirmation": True,
        "metadata": {
            "created_at": utc_now_iso(),
        },
    }


def send_email_via_smtp(email_draft: dict[str, Any]) -> dict[str, Any]:
    """
    Send email via SMTP if configured.

    Required .env:
        SMTP_HOST=
        SMTP_PORT=587
        SMTP_USERNAME=
        SMTP_PASSWORD=
        SMTP_FROM=  # optional, defaults to SMTP_USERNAME

    Returns a result dict. It never raises to the graph.
    """
    smtp_host = os.getenv("SMTP_HOST")
    smtp_port = int(os.getenv("SMTP_PORT", "587"))
    smtp_username = os.getenv("SMTP_USERNAME")
    smtp_password = os.getenv("SMTP_PASSWORD")
    smtp_from = os.getenv("SMTP_FROM") or smtp_username

    if not all([smtp_host, smtp_username, smtp_password, smtp_from]):
        return {
            "sent": False,
            "reason": "SMTP is not configured in .env",
        }

    try:
        msg = EmailMessage()
        msg["From"] = smtp_from
        msg["To"] = email_draft["to"]
        msg["Subject"] = email_draft["subject"]
        msg.set_content(email_draft.get("body") or "")

        attachment_path = email_draft.get("attachment_path")
        if attachment_path and Path(attachment_path).exists():
            path = Path(attachment_path)
            msg.add_attachment(
                path.read_bytes(),
                maintype="application",
                subtype="octet-stream",
                filename=path.name,
            )

        # Important: this must be OUTSIDE the attachment block.
        # Otherwise emails without attachments are never sent.
        with smtplib.SMTP(smtp_host, smtp_port, timeout=20) as server:
            server.ehlo()
            server.starttls()
            server.ehlo()
            server.login(smtp_username, smtp_password)
            server.send_message(msg)

        return {
            "sent": True,
            "to": email_draft["to"],
            "subject": email_draft["subject"],
            "sent_at": utc_now_iso(),
        }

    except smtplib.SMTPAuthenticationError as exc:
        return {
            "sent": False,
            "reason": "SMTP authentication failed. Check SMTP_USERNAME and SMTP_PASSWORD / app password.",
        }

    except smtplib.SMTPConnectError as exc:
        return {
            "sent": False,
            "reason": "SMTP connection failed. Check SMTP_HOST and SMTP_PORT.",
        }

    except TimeoutError as exc:
        return {
            "sent": False,
            "reason": "SMTP timed out after 20 seconds. The mail server did not respond fast enough.",
        }

    except OSError as exc:
        return {
            "sent": False,
            "reason": f"SMTP network error: {exc}",
        }

    except Exception as exc:
        return {
            "sent": False,
            "reason": str(exc),
        }


# ============================================================
# RAG / DOCUMENT RETRIEVAL
# ============================================================




# ============================================================
# COMMAND / CODE ARTIFACT HELPERS
# ============================================================

def save_commands_artifact(commands_text: str, filename: Optional[str] = None) -> dict[str, Any]:
    if filename is None:
        filename = f"{new_id('commands')}.md"

    if not filename.lower().endswith(".md"):
        filename += ".md"

    return save_markdown(
        commands_text,
        filename=filename,
        title="Generated Commands",
        artifact_type="commands",
        created_from="command_code_agent",
    )


def save_report_artifact(report_text: str, filename: Optional[str] = None) -> dict[str, Any]:
    if filename is None:
        filename = f"{new_id('report')}.md"

    if not filename.lower().endswith(".md"):
        filename += ".md"

    return save_markdown(
        report_text,
        filename=filename,
        title="Generated Report",
        artifact_type="report",
        created_from="report_agent",
    )


# ============================================================
# PROMPT FORMATTERS
# ============================================================

def format_recent_messages(messages: list[dict[str, Any]], limit: int = 8) -> str:
    if not messages:
        return "No recent messages."

    selected = messages[-limit:]
    lines: list[str] = []

    for msg in selected:
        role = msg.get("role", "unknown")
        content = msg.get("content", "")
        lines.append(f"{role}: {short_preview(content, 1000)}")

    return "\n".join(lines)


def format_artifacts_for_prompt(artifacts: list[dict[str, Any]], limit: int = 6) -> str:
    if not artifacts:
        return "No artifacts."

    lines: list[str] = []

    for artifact in artifacts[-limit:]:
        lines.append(
            f"- id={artifact.get('artifact_id')} | "
            f"type={artifact.get('type')} | "
            f"title={artifact.get('title')} | "
            f"path={artifact.get('path')} | "
            f"preview={short_preview(artifact.get('content_preview', ''), 250)}"
        )

    return "\n".join(lines)


def format_state_summary_for_prompt(state: dict[str, Any]) -> str:
    """
    Compact state summary for orchestrator/context resolver prompts.
    """
    return (
        f"current_topic: {state.get('current_topic')}\n"
        f"active_task: {state.get('active_task')}\n"
        f"pending_approval: {state.get('pending_approval')}\n"
        f"latest_answer: {short_preview(state.get('latest_answer') or '', 400)}\n"
        f"latest_report_id: {state.get('latest_report_id')}\n"
        f"latest_artifact_id: {state.get('latest_artifact_id')}\n"
        f"latest_code_id: {state.get('latest_code_id')}\n"
        f"latest_email_draft_id: {state.get('latest_email_draft_id')}\n"
        f"recent_messages:\n{format_recent_messages(state.get('messages', []))}\n"
        f"artifacts:\n{format_artifacts_for_prompt(state.get('artifacts', []))}"
    )
from pathlib import Path
import re
from typing import Optional

SUPPORTED_EXPORT_EXTENSIONS = {
    "pdf": ".pdf",
    "markdown": ".md",
    "md": ".md",
    "text": ".txt",
    "txt": ".txt",
    "word": ".docx",
    "docx": ".docx",
}

SENSITIVE_FILENAME_WORDS = [
    "secret", "key", "api", "token", "password", "credential",
    "email", "gmail", "private", "logs", "vulnerable"
]


def compact_artifact_basename(
    requested_filename: str = "",
    *,
    fallback: str = "artifact",
    max_words: int = 4,
    max_chars: int = 42,
) -> str:
    """
    Make a short, safe, non-sensitive base filename.
    No extension here.
    """
    raw = requested_filename or fallback or "artifact"

    # remove extension if model/user included it
    raw = Path(raw).stem

    raw = raw.lower()
    raw = re.sub(r"[^a-z0-9\u0600-\u06FF]+", "_", raw, flags=re.UNICODE)
    raw = re.sub(r"_+", "_", raw).strip("_")

    words = [w for w in raw.split("_") if w]

    # remove sensitive words from filename
    words = [w for w in words if w not in SENSITIVE_FILENAME_WORDS]

    if not words:
        words = [fallback]

    # keep it short
    base = "_".join(words[:max_words])
    base = base[:max_chars].strip("_")

    return base or fallback or "artifact"


def ensure_export_extension(filename: str, export_format: str) -> str:
    """
    Force the filename extension to match requested format.
    """
    ext = SUPPORTED_EXPORT_EXTENSIONS.get((export_format or "").lower(), ".txt")

    base = Path(filename).stem
    base = base or "artifact"

    return f"{base}{ext}"


def build_export_filename(
    requested_filename: str = "",
    *,
    fallback: str = "artifact",
    export_format: str = "text",
) -> str:
    base = compact_artifact_basename(
        requested_filename,
        fallback=fallback,
    )
    return ensure_export_extension(base, export_format)
# ============================================================
# RAG RE-EXPORTS
# ============================================================

from my_agent.rag import (
    build_index_from_folder,
    retrieve_from_documents,
    rag_status,
    clear_vectorstore,
    is_index_valid,
)


def ingest_pdfs_to_vectorstore(folder_path: Optional[str] = None) -> dict[str, Any]:
    """
    Backwards-compatible alias for older code that used this name.
    """
    if folder_path:
        return build_index_from_folder(folder_path)
    return build_index_from_folder()
