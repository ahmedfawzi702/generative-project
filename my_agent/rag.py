# my_agent/rag.py
"""
RAG pipeline:
- Load files from FILES_DIR
- Extract text from PDF/DOCX/CSV/TXT/MD/JSON
- Split into chunks
- Embed chunks locally with Ollama nomic-embed-text
- Store vectors in local Qdrant
- Retrieve relevant chunks for RAG Agent

This file does NOT generate final answers.
nodes.py handles answer generation using OpenRouter LLMs.
"""

from __future__ import annotations

import hashlib
import json
import os
from functools import lru_cache
from pathlib import Path
from typing import Any, Dict, List, Optional

import fitz
import pandas as pd
from docx import Document as DocxDocument
from dotenv import load_dotenv
from langchain_core.documents import Document
from langchain_ollama import OllamaEmbeddings
from langchain_qdrant import QdrantVectorStore
from langchain_text_splitters import RecursiveCharacterTextSplitter
from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, VectorParams

load_dotenv()


# ============================================================
# CONFIG
# ============================================================

FILES_DIR = os.getenv("FILES_DIR", os.getenv("RAW_PDFS_DIR", "data/raw_pdfs"))
QDRANT_PATH = os.getenv("QDRANT_PATH", "data/vectorstore/qdrant")
META_FILE = os.getenv("RAG_META_FILE", "data/vectorstore/rag_meta.json")

COLLECTION_NAME = os.getenv("QDRANT_COLLECTION_NAME", "cybersecurity_docs")

OLLAMA_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
EMBED_MODEL = os.getenv("EMBEDDING_MODEL_NAME", "nomic-embed-text:latest")
EMBED_DIM = int(os.getenv("EMBEDDING_DIM", "768"))

CHUNK_SIZE = int(os.getenv("RAG_CHUNK_SIZE", "900"))
CHUNK_OVERLAP = int(os.getenv("RAG_CHUNK_OVERLAP", "120"))

DEFAULT_K = int(os.getenv("RAG_RETRIEVAL_K", "4"))
DEFAULT_FETCH_K = int(os.getenv("RAG_FETCH_K", "10"))

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".csv", ".txt", ".md", ".json"}


# ============================================================
# BASIC HELPERS
# ============================================================

def ensure_dirs() -> None:
    Path(QDRANT_PATH).mkdir(parents=True, exist_ok=True)
    Path(META_FILE).parent.mkdir(parents=True, exist_ok=True)
    Path(FILES_DIR).mkdir(parents=True, exist_ok=True)


def normalize_text(text: str) -> str:
    return "\n".join(line.strip() for line in text.splitlines() if line.strip()).strip()


def file_sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def save_meta(meta: Dict[str, Any]) -> None:
    ensure_dirs()
    with open(META_FILE, "w", encoding="utf-8") as f:
        json.dump(meta, f, indent=2, ensure_ascii=False)


def load_meta() -> Dict[str, Any]:
    if not os.path.exists(META_FILE):
        return {}

    try:
        with open(META_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def get_all_files_from_folder(folder_path: str = FILES_DIR) -> List[Path]:
    folder = Path(folder_path)

    if not folder.exists():
        return []

    return [
        p for p in folder.iterdir()
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
    ]


# ============================================================
# OLLAMA EMBEDDINGS + QDRANT
# ============================================================

def get_qdrant_client() -> QdrantClient:
    ensure_dirs()
    return QdrantClient(path=QDRANT_PATH)


@lru_cache(maxsize=1)
def get_embeddings() -> OllamaEmbeddings:
    """
    Local embeddings through Ollama.

    Make sure:
        ollama serve
        ollama list
    includes:
        nomic-embed-text:latest
    """
    return OllamaEmbeddings(
        model=EMBED_MODEL,
        base_url=OLLAMA_URL,
    )


def collection_exists(client: QdrantClient, collection_name: str) -> bool:
    try:
        response = client.get_collections()

        if hasattr(response, "collections"):
            collections = response.collections
        elif hasattr(response, "result") and hasattr(response.result, "collections"):
            collections = response.result.collections
        elif isinstance(response, dict):
            if "collections" in response:
                collections = response["collections"]
            elif "result" in response and "collections" in response["result"]:
                collections = response["result"]["collections"]
            else:
                collections = []
        else:
            collections = []

        names = []

        for item in collections:
            if hasattr(item, "name"):
                names.append(item.name)
            elif isinstance(item, dict) and "name" in item:
                names.append(item["name"])

        return collection_name in names

    except Exception:
        return False


def is_index_valid() -> bool:
    meta = load_meta()

    if not meta.get("index_signature"):
        return False

    if meta.get("files_count", 0) <= 0:
        return False

    if meta.get("chunks_count", 0) <= 0:
        return False

    try:
        client = get_qdrant_client()
        return collection_exists(client, COLLECTION_NAME)
    except Exception:
        return False


# ============================================================
# DOCUMENT EXTRACTION
# ============================================================

def extract_pdf(file_path: Path) -> List[Document]:
    docs: List[Document] = []

    pdf = fitz.open(file_path)

    try:
        for page_num, page in enumerate(pdf, start=1):
            text = normalize_text(page.get_text("text"))

            if text:
                docs.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": file_path.name,
                            "source_name": file_path.name,
                            "source_path": str(file_path),
                            "page": page_num,
                            "file_type": "pdf",
                        },
                    )
                )
    finally:
        pdf.close()

    return docs


def extract_docx(file_path: Path) -> List[Document]:
    doc = DocxDocument(str(file_path))

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    table_rows = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            )
            if row_text:
                table_rows.append(row_text)

    text = "\n".join(paragraphs + table_rows).strip()

    return [
        Document(
            page_content=text if text else "[Empty DOCX]",
            metadata={
                "source": file_path.name,
                "source_name": file_path.name,
                "source_path": str(file_path),
                "page": None,
                "file_type": "docx",
            },
        )
    ]


def extract_csv(file_path: Path) -> List[Document]:
    try:
        df = pd.read_csv(file_path)
        text = df.astype(str).to_csv(index=False)
    except Exception:
        text = file_path.read_text(encoding="utf-8", errors="ignore")

    return [
        Document(
            page_content=text if text.strip() else "[Empty CSV]",
            metadata={
                "source": file_path.name,
                "source_name": file_path.name,
                "source_path": str(file_path),
                "page": None,
                "file_type": "csv",
            },
        )
    ]


def extract_text_file(file_path: Path) -> List[Document]:
    text = file_path.read_text(encoding="utf-8", errors="ignore")

    return [
        Document(
            page_content=text if text.strip() else "[Empty Text File]",
            metadata={
                "source": file_path.name,
                "source_name": file_path.name,
                "source_path": str(file_path),
                "page": None,
                "file_type": file_path.suffix.lower().replace(".", ""),
            },
        )
    ]


def extract_documents(file_path: Path) -> List[Document]:
    ext = file_path.suffix.lower()

    if ext == ".pdf":
        return extract_pdf(file_path)

    if ext == ".docx":
        return extract_docx(file_path)

    if ext == ".csv":
        return extract_csv(file_path)

    if ext in {".txt", ".md", ".json"}:
        return extract_text_file(file_path)

    return []


# ============================================================
# SPLITTING + BUILDING INDEX
# ============================================================

def split_documents(documents: List[Document]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    return splitter.split_documents(documents)


def build_vectorstore(chunks: List[Document]) -> None:
    """
    Create/rebuild Qdrant collection and add chunks.
    """
    client = get_qdrant_client()

    if collection_exists(client, COLLECTION_NAME):
        client.delete_collection(COLLECTION_NAME)

    client.create_collection(
        collection_name=COLLECTION_NAME,
        vectors_config=VectorParams(size=EMBED_DIM, distance=Distance.COSINE),
    )

    vectorstore = QdrantVectorStore(
        client=client,
        collection_name=COLLECTION_NAME,
        embedding=get_embeddings(),
    )

    batch_size = 100

    for i in range(0, len(chunks), batch_size):
        vectorstore.add_documents(chunks[i:i + batch_size])


def build_index_from_folder(folder_path: str = FILES_DIR) -> Dict[str, Any]:
    """
    Full indexing pipeline.

    Files
    -> extract text
    -> split chunks
    -> embed chunks with Ollama
    -> store vectors in Qdrant
    -> save meta
    """
    ensure_dirs()

    files = get_all_files_from_folder(folder_path)

    if not files:
        raise ValueError(f"No supported files found in: {folder_path}")

    all_docs: List[Document] = []
    file_info = []

    for file_path in files:
        file_bytes = file_path.read_bytes()
        docs = extract_documents(file_path)

        all_docs.extend(docs)

        file_info.append(
            {
                "name": file_path.name,
                "path": str(file_path),
                "size_bytes": len(file_bytes),
                "sha256": file_sha256(file_bytes),
            }
        )

    if not all_docs:
        raise ValueError("Files found but no readable text could be extracted.")

    chunks = split_documents(all_docs)

    if not chunks:
        raise ValueError("Documents extracted but no chunks were created.")

    build_vectorstore(chunks)
    get_cached_vectorstore.cache_clear()

    index_signature = hashlib.sha256(
        json.dumps(file_info, sort_keys=True).encode("utf-8")
    ).hexdigest()

    meta = {
        "folder": folder_path,
        "files_count": len(files),
        "pages_or_docs_count": len(all_docs),
        "chunks_count": len(chunks),
        "embedding_model": EMBED_MODEL,
        "embedding_dim": EMBED_DIM,
        "ollama_base_url": OLLAMA_URL,
        "collection_name": COLLECTION_NAME,
        "qdrant_path": QDRANT_PATH,
        "chunk_size": CHUNK_SIZE,
        "chunk_overlap": CHUNK_OVERLAP,
        "index_signature": index_signature,
        "files": file_info,
    }

    save_meta(meta)
    return meta


# Alias عشان لو عايز تستخدم نفس الاسم اللي اتكلمنا عنه قبل كده.
ingest_pdfs_to_vectorstore = build_index_from_folder


# ============================================================
# LOADING + RETRIEVAL
# ============================================================

def load_vectorstore() -> QdrantVectorStore:
    client = get_qdrant_client()

    return QdrantVectorStore(
        client=client,
        collection_name=COLLECTION_NAME,
        embedding=get_embeddings(),
    )


@lru_cache(maxsize=1)
def get_cached_vectorstore() -> QdrantVectorStore:
    return load_vectorstore()


def get_active_vectorstore() -> Optional[QdrantVectorStore]:
    if not is_index_valid():
        return None

    try:
        return get_cached_vectorstore()
    except Exception:
        get_cached_vectorstore.cache_clear()
        return None


def retrieve_documents(
    query: str,
    *,
    k: int = DEFAULT_K,
    fetch_k: int = DEFAULT_FETCH_K,
) -> List[Document]:
    vectorstore = get_active_vectorstore()

    if vectorstore is None:
        return []

    retriever = vectorstore.as_retriever(
        search_type="mmr",
        search_kwargs={
            "k": k,
            "fetch_k": fetch_k,
        },
    )

    return retriever.invoke(query)


def retrieve_from_documents(query: str, k: int = DEFAULT_K) -> List[Dict[str, Any]]:
    """
    Function used by nodes.py rag_agent.

    Query embedding happens automatically inside QdrantVectorStore/OllamaEmbeddings.
    """
    if not query or not query.strip():
        return []

    try:
        docs = retrieve_documents(query, k=k, fetch_k=DEFAULT_FETCH_K)

        results = []

        for doc in docs:
            metadata = doc.metadata or {}

            source_name = (
                metadata.get("source_name")
                or metadata.get("source")
                or "unknown"
            )
            source_path = (
                metadata.get("source_path")
                or metadata.get("source")
                or ""
            )
            page = metadata.get("page")
            file_type = metadata.get("file_type", "document")

            title = source_name
            if page:
                title = f"{source_name} - page {page}"

            results.append(
                {
                    "source_id": f"rag_{hashlib.sha1((source_name + str(page) + doc.page_content[:50]).encode('utf-8')).hexdigest()[:10]}",
                    "type": file_type,
                    "title": title,
                    "url": None,
                    "path": source_path,
                    "content": doc.page_content,
                    "snippet": doc.page_content[:500],
                    "score": None,
                    "metadata": {
                        **metadata,
                        "query": query,
                        "embedding_model": EMBED_MODEL,
                    },
                }
            )

        return results

    except Exception as exc:
        return [
            {
                "source_id": "rag_error",
                "type": "tool",
                "title": "RAG retrieval error",
                "url": None,
                "path": "",
                "content": str(exc),
                "snippet": str(exc),
                "score": 0.0,
                "metadata": {
                    "error": True,
                    "query": query,
                    "embedding_model": EMBED_MODEL,
                },
            }
        ]


# ============================================================
# STATUS + CLEAR
# ============================================================

def rag_status() -> Dict[str, Any]:
    meta = load_meta()
    ready = is_index_valid()

    return {
        "index_ready": ready,
        "meta": meta,
        "files_dir": FILES_DIR,
        "qdrant_path": QDRANT_PATH,
        "collection_name": COLLECTION_NAME,
        "embedding_model": EMBED_MODEL,
        "ollama_base_url": OLLAMA_URL,
    }


def clear_vectorstore() -> Dict[str, Any]:
    try:
        client = get_qdrant_client()

        if collection_exists(client, COLLECTION_NAME):
            client.delete_collection(COLLECTION_NAME)

        if os.path.exists(META_FILE):
            os.remove(META_FILE)

        get_cached_vectorstore.cache_clear()

        return {
            "success": True,
            "message": "RAG vectorstore and metadata cleared.",
            "collection_name": COLLECTION_NAME,
        }

    except Exception as exc:
        return {
            "success": False,
            "message": str(exc),
        }